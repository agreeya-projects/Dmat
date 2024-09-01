from flask import Flask, request, jsonify, send_file
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain.chains import ConversationChain
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from flask_cors import CORS
from langchain_openai import OpenAI
import tiktoken
import json
from langchain_openai import ChatOpenAI
from langchain_mongodb import MongoDBAtlasVectorSearch
from openai import OpenAI as openai_OpenAI
import urllib.parse
import os
from dotenv import load_dotenv
from pymongo import MongoClient
from utils import set_open_key
from langchain_community.embeddings import HuggingFaceEmbeddings
from docx import Document as DocxDocument

# Load environment variables from .env file
load_dotenv()

# Set OpenAI API key
set_open_key()

# MongoDB credentials
username = urllib.parse.quote_plus(os.getenv('MONGO_USERNAME'))
password = urllib.parse.quote_plus(os.getenv('MONGO_PASSWORD'))
atlas_connection_string = f"mongodb+srv://{username}:{password}@cluster0.qualb0u.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"

# Connect to MongoDB
mongodb_client = MongoClient(atlas_connection_string)
print("-------------------------------------- MongoDB Connected ---------------------------------------------")
print(mongodb_client.list_database_names())

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

client = openai_OpenAI()
model_name = "gpt-3.5-turbo"
encoding_model_name = "gpt-3.5-turbo"
collection_name_suffix = "dmat_doc"

instructions = {
    "Introduction": "For Introduction of the DMAT tool, generate a detailed paragraph highlighting its main features, functionalities, and use cases. Mention how it provides real-time monitoring and analysis for Android 5G NR and LTE devices. Include its compatibility, logging capabilities, Single-Sign-On authentication, and remote control features. Make the introduction informative and engaging for potential users and stakeholders. Please include any additional relevant information to make the introduction comprehensive and detailed.",
    "Purpose of the SDD": "Write a comprehensive and detailed paragraph for the 'Purpose of the SDD' section. Explain the objective and significance of the Software Design Document (SDD) in the context of the DMAT tool. Describe how the SDD serves as a comprehensive blueprint for the software architecture, detailing specific design decisions and implementation strategies. Emphasize the importance of the SDD as a critical reference document for developers, testers, and stakeholders, ensuring alignment and coherence throughout the project lifecycle. Discuss how the SDD facilitates effective communication and understanding among all project participants, helping to identify and mitigate potential risks and challenges early in the process. Highlight the role of the SDD in maintaining consistency and clarity, ensuring that all aspects of the project are thoroughly documented and understood. Explain how the SDD supports continuous improvement and future scalability of the DMAT tool. Conclude by discussing how the SDD is indispensable for the successful implementation, deployment, and maintenance of the DMAT tool, ultimately contributing to its long-term success and efficiency. Please provide additional relevant details to ensure thorough coverage of the purpose.",
    "General Overview": "Write a comprehensive and detailed paragraph for the 'General Overview' section of the DMAT tool. Describe the primary purpose of DMAT, highlighting its role in providing real-time monitoring and analysis for Android 5G NR and LTE devices. Explain how DMAT supports efficient device performance monitoring and facilitates various testing environments. Detail its advanced logging capabilities, emphasizing how they contribute to effective data collection and analysis. Discuss the tool's compatibility with a wide range of devices and the benefits of this broad compatibility. Highlight the Single-Sign-On authentication feature, explaining its role in enhancing security and simplifying user management. Describe the remote control features, including how they allow users to manage multiple devices simultaneously and improve operational efficiency. Mention the integration of DMAT with existing post-processing tools and the advantages this offers to Verizon's testing teams. Conclude by explaining how DMAT enhances overall productivity and efficiency, providing a robust and cost-effective solution for comprehensive device performance monitoring and analysis. Please include any additional relevant information to provide a thorough overview.",
    "Key Pain Points Addressed": "Write a comprehensive and detailed paragraph for the 'Key Pain Points Addressed' section. Identify and describe each key pain point that the DMAT tool addresses. Include specific challenges related to real-time monitoring and analysis of Android 5G NR and LTE devices. Discuss how DMAT resolves issues such as inconsistent data collection, difficulty in managing multiple devices, and security concerns. Highlight the tool's advanced logging capabilities and how they improve data accuracy and reliability. Explain how the Single-Sign-On authentication feature enhances security and simplifies user management. Detail how the remote control features help in efficiently managing multiple devices simultaneously, reducing operational complexity. Mention the integration with existing post-processing tools and how it streamlines the analysis process for testing teams. Conclude by summarizing how DMAT effectively addresses these pain points, leading to improved efficiency, productivity, and overall performance in device monitoring and analysis. Please provide additional relevant details to ensure thorough coverage of key pain points.",
    "Use Case Solution": "Generate a detailed and comprehensive paragraph for the 'Use Case Solution' section. Describe specific scenarios where the DMAT tool is effectively used, highlighting real-time monitoring and analysis, data collection improvements, secure access via Single-Sign-On, remote management of multiple devices, and integration with post-processing tools. Provide examples of how DMAT enhances operational efficiency and reduces complexity in real-world applications.",
    "Goals and Expected Outcomes": "Generate a detailed and comprehensive paragraph for the 'Goals and Expected Outcomes' section. Outline the primary goals of implementing the DMAT tool, such as improving real-time monitoring and analysis, enhancing data accuracy, and streamlining device management. Describe the expected outcomes, including increased operational efficiency, better security through Single-Sign-On, simplified user management, and more effective data integration with post-processing tools. Highlight the benefits to users and stakeholders, leading to improved performance and productivity.",
    "Scope of Proof of Concept (POC)": "Generate a detailed and comprehensive paragraph for the 'Scope of Proof of Concept (POC)' section. Define the objectives and boundaries of the POC for the DMAT tool. Describe the specific features and functionalities to be tested, including real-time monitoring, advanced logging capabilities, Single-Sign-On authentication, and remote device management. Explain the criteria for success, performance benchmarks, user feedback, and integration with existing systems. Highlight the key deliverables and timeline for completing the POC, and discuss how the POC outcomes will inform the full-scale implementation and potential improvements for the DMAT tool.",
    "Assumptions": "List and explain the assumptions made during the development of the DMAT tool.",
    "Technical and Operational Assumptions": "Describe the technical and operational assumptions that underlie the DMAT tool's design and functionality.",
    "Project Management and Collaboration Assumptions": "Outline the project management and collaboration assumptions made for the successful development and deployment of the DMAT tool.",
    "Integration Assumptions": "Detail the integration assumptions necessary for the DMAT tool to function within the existing systems.",
    "System Integration": "Discuss the integration processes and procedures for the DMAT tool within the existing systems.",
    "Dependencies": "Identify the dependencies that the DMAT tool relies on for successful operation.",
    "Technical Dependencies": "Explain the technical dependencies required for the DMAT tool.",
    "Resource Dependencies": "Outline the resource dependencies for the DMAT tool.",
    "Operational Dependencies": "Describe the operational dependencies necessary for the DMAT tool's functionality.",
    "Real-Time Assistance": "Detail the real-time assistance features available in the DMAT tool.",
    "User Authentication System": "Describe the user authentication system used in the DMAT tool.",
    "Approval and Documentation Dependencies": "Discuss the approval processes and documentation dependencies for the DMAT tool.",
    "Approval Processes": "Outline the approval processes required for the DMAT tool.",
    "Documentation and Specifications": "Detail the documentation and specifications needed for the DMAT tool.",
    "Limitations / Out of Scope": "Identify the limitations and out-of-scope elements of the DMAT tool.",
    "Data and Content Limitations": "Explain the data and content limitations of the DMAT tool.",
    "Exploration Beyond Confluence": "Detail the limitations regarding exploration beyond Confluence.",
    "Contextual Searches in Visual Elements": "Describe the limitations related to contextual searches in visual elements.",
    "Analysis and Extraction in Complex Documents": "Explain the limitations in analyzing and extracting content from complex documents.",
    "Text Searches Within Images": "Detail the limitations of text searches within images.",
    "System and Technical Limitations": "Discuss the system and technical limitations of the DMAT tool.",
    "Ticket Management System Changes": "Outline the limitations related to ticket management system changes.",
    "Network and Third-Party Service Issues": "Identify the limitations concerning network and third-party service issues.",
    "Searches and Third-Party Service Issues": "Detail the limitations regarding searches and third-party service issues.",
    "High Level Architecture and Design": "Provide a high-level overview of the architecture and design of the DMAT tool.",
    "System Architecture and Design": "Describe the system architecture and design of the DMAT tool.",
    "High Level Frontend Class Diagram": "Provide a high-level frontend class diagram for the DMAT tool.",
    "Technologies Used with Version": "List and explain the technologies used in the DMAT tool along with their versions.",
    "Security Procedures": "Detail the security procedures implemented in the DMAT tool.",
    "Information Architecture – Data Flow": "Describe the information architecture and data flow of the DMAT tool.",
    "Data Flow Diagrams (DFD)": "Provide data flow diagrams (DFD) for the DMAT tool.",
    "Entity-Relationship Diagrams (ERD)": "Provide entity-relationship diagrams (ERD) for the DMAT tool.",
    "Design and Usability": "Discuss the design and usability aspects of the DMAT tool.",
    "Key Features of the Project UI": "Describe the key features of the DMAT tool's user interface.",
    "Application Data Flow Class Diagram": "Provide an application data flow class diagram for the DMAT tool.",
    "Key Components of the Application Data Flow": "Identify the key components of the application data flow in the DMAT tool.",
    "System Design": "Discuss the detailed system design of the DMAT tool.",
    "Detailed Requirements and User Stories": "Provide detailed requirements and user stories for the DMAT tool.",
    "Functional Requirements": "List and explain the functional requirements of the DMAT tool.",
    "Non-Functional Requirements": "List and explain the non-functional requirements of the DMAT tool.",
    "User Story Mapping": "Provide a user story mapping for the DMAT tool.",
    "Database Design": "Discuss the database design of the DMAT tool.",
    "Data Objects and Resultant Data Structures": "Detail the data objects and resultant data structures of the DMAT tool.",
    "Schema Definitions and Table Relationships": "Provide schema definitions and table relationships for the DMAT tool.",
    "Detailed Design": "Describe the detailed design of the DMAT tool.",
    "Software Detailed Design": "Provide a detailed software design for the DMAT tool.",
    "Module Descriptions": "Describe the various modules of the DMAT tool.",
    "Algorithm Specifications": "Detail the algorithm specifications used in the DMAT tool.",
    "Error Handling and Logging Mechanisms": "Explain the error handling and logging mechanisms in the DMAT tool.",
    "Configuration Management": "Discuss the configuration management processes for the DMAT tool.",
    "External Interfaces": "Identify the external interfaces used by the DMAT tool.",
    "Interface & API References & Details": "Provide references and details for the interfaces and APIs used in the DMAT tool.",
    "API Endpoints and Specifications": "Detail the API endpoints and specifications for the DMAT tool.",
    "Authentication and Authorization Mechanisms": "Explain the authentication and authorization mechanisms used in the DMAT tool.",
    "Error Codes and Responses": "Provide the error codes and responses used in the DMAT tool.",
    "FAQs": "List common questions and answers related to the DMAT tool.",
    "Troubleshooting Guide": "Provide a troubleshooting guide for the DMAT tool.",
    "User Stories": "List and describe the comprehensive user stories for the DMAT tool.",
    "Acceptance Criteria": "Provide the acceptance criteria for the user stories of the DMAT tool.",
    "Testing and Validation": "Discuss the testing and validation processes for the DMAT tool.",
    "Deployment and Maintenance": "Discuss the deployment and maintenance plans for the DMAT tool.",
    "Deployment Plan": "Provide the deployment plan for the DMAT tool.",
    "Rollback Procedures": "Detail the rollback procedures for the DMAT tool.",
    "Maintenance Plan": "Explain the maintenance plan for the DMAT tool.",
    "Monitoring and Alerting": "Describe the monitoring and alerting mechanisms for the DMAT tool.",
    "Backup and Recovery": "Provide the backup and recovery processes for the DMAT tool.",
    "Appendix": "Include additional information related to the DMAT tool.",
    "Record of Changes": "Maintain a record of changes for the DMAT tool documentation.",
    "Acronyms": "List and explain the acronyms used in the DMAT tool documentation.",
    "Referenced Documents": "Provide references to documents related to the DMAT tool.",
    "Approvals": "List the approvals required for the DMAT tool documentation."
}

PROMPT_TEMPLATE = """
Generate detailed content for the {section} section of the documentation of the tool named {tool_name} based on the following instructions:
{instruction}

Use information from the database to enrich the content. Ensure that each section is comprehensive and detailed. 
Avoid generic descriptions and focus on specifics related to the tool.

The output content should not contain any extra text such as "as asked in query", "as an AI assistant", etc.
The output content should be in such a way that it can directly be added to a doc file using python code.
The output content should follow the given template and include all sections and subsections.

The output text should be like so that its can be directly converted in doc file using python.

Give the response in professional way.
"""

def generate_documentation_prompt(section, instruction, tool_name, context_text):
    prompt = PROMPT_TEMPLATE.format(
        section=section,
        instruction=instruction,
        tool_name=tool_name,
    )
    # print("------------------- Prompt------------------")
    # print(prompt)
    # print("-------------------Final prompt-------------")
    if context_text:
        prompt += "\n\nContext:\n" + context_text
        print(prompt)
    return prompt


bos = "<|begin_of_text|><|start_header_id|>system<|end_header_id|>\n"
DEFAULT_SYSTEM_PROMPT='''\nYou are a helpful, smart, kind, and efficient AI assistant. 
You are a professional tool that generates documentation of any tool or script using the data provided in the database..'''

user = "\n<|eot_id|><|start_header_id|>user<|end_header_id|>\n"

assistant = "\n<|eot_id|><|start_header_id|>assistant<|end_header_id|>"

SYSTEM_PROMPT = bos + DEFAULT_SYSTEM_PROMPT + user

# B_INST, E_INST = "[INST]", "[/INST]"
# B_SYS, E_SYS = "<>\n", "\n<>\n\n"

# DEFAULT_SYSTEM_PROMPT = """
# You are a professional tool that generates documentation of any tool or script using the data provided in the database.
# """

# SYSTEM_PROMPT = B_SYS + DEFAULT_SYSTEM_PROMPT + E_SYS

db_name = "DMAT"
model_selected = "llama3"
max_tokens = 4000  # Set max tokens limit

if model_selected == "openai":
    local_embedding = False
    collection_name = "openai"
    llm = ChatOpenAI(model=model_name, max_tokens=max_tokens)
    llm_agent = openai_OpenAI
elif model_selected == "llama3":
    from langchain_community.llms import CTransformers
    from langchain.llms import LlamaCpp
    local_embedding = True
    collection_name = "Information_Store"
    llm = LlamaCpp(
                #streaming = True,
                model_path="models/Meta-Llama-3-8B-Instruct.Q4_K_M.gguf",
                # config={#'max_new_tokens':128,
                #         'temperature':0.01,
                #         'lctx': 1500}
                temperature=0.8,
                # top_p=1,
                # verbose=True,
                n_ctx=4000
                )
    llm_agent = llm

chat_history = []
encoding = tiktoken.encoding_for_model(encoding_model_name)

memory = ConversationBufferWindowMemory(llm=llm_agent, max_token_limit=4000)
conversation = ConversationChain(llm=llm, memory=memory, verbose=True)

def rephrase(user_query):
    rephrase_prompt = f"""Please analyze the given text for grammatical errors and provide corrections. Focus on punctuation, sentence structure, and overall clarity to ensure a polished and error-free piece. 
    Input text: {user_query}
    """

    completion = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": rephrase_prompt}
        ]
    )

    print(completion.choices[0].message.content)
    try:
        response_content = completion.choices[0].message.content
        json_start = response_content.find("{")
        json_end = response_content.rfind("}") + 1
        rephrased_answer = json.loads(response_content[json_start:json_end])["answer"]
    except (ValueError, KeyError):
        rephrased_answer = completion.choices[0].message.content
    return rephrased_answer

def get_vector_index_mongo_for_lang(collection_name, atlas_connection_string, db_name): #collection_name_suffix, 
    if local_embedding:
        embed_model = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5")  # Use HuggingFaceEmbeddings
    else:
        embed_model = OpenAIEmbeddings()
    #collection_name = collection_name + "_" + collection_name_suffix
    vector_search = MongoDBAtlasVectorSearch.from_connection_string(
        atlas_connection_string,
        f"{db_name}.{collection_name}",
        embed_model,
        index_name=collection_name + "_index"
    )
    return vector_search

db = get_vector_index_mongo_for_lang(collection_name, atlas_connection_string, db_name) # collection_name_suffix,

@app.route('/query', methods=['POST'])
def query():
    global chat_history
    global db

    tool_name = request.form.get('tool_name', 'DMAT')
    sections = [
        "Introduction",
        "Purpose of the SDD",
        "General Overview",
        # "Key Pain Points Addressed",
        # "Use Case Solution",
        # "Goals and Expected Outcomes",
        # "Scope of Proof of Concept (POC)",
        # "Assumptions",
        # "System Integration",
        # "Limitations / Out of Scope",
        # "High Level Architecture and Design",
        # "Security Procedures",
        # "Information Architecture – Data Flow",
        # "Design and Usability",
        # "System Design",
        # "External Interfaces",
        # "FAQs",
        # "User Stories",
        # "Testing and Validation",
        # "Deployment and Maintenance",
        # "Appendix",
    ]

    document = DocxDocument()
    document.add_heading('Documentation', 0)



    for section in sections:
        instruction = instructions.get(section, "")
        query_text = f"Generate detailed content for the {section} section."
        context_text = "\n".join([doc.page_content for doc, score in db.similarity_search_with_score(query_text, k=3)])
        print('Context_text:- ', context_text)
        
        if model_selected == "openai":
            query_text = rephrase(query_text)
        
        
        prompt = generate_documentation_prompt(section, instruction, tool_name, context_text)
        prompt_template = ChatPromptTemplate.from_template(prompt)
        #print(prompt_template)
        
        # if len(chat_history) > 5:
        #     chat_history = chat_history[-5:]
        
        formatted_prompt = prompt_template.format(context=context_text) # , history=chat_history
        
        if model_selected == "llama3":
            #formatted_prompt = SYSTEM_PROMPT + prompt + assistant + "\n" + (chat_history[-1] if chat_history else "")

            formatted_prompt = SYSTEM_PROMPT + formatted_prompt + assistant
        
        response_text = conversation.predict(input=formatted_prompt)
        #chat_history.append(f"User: {query_text} | Bot: {response_text}")
        
        formatted_response_text = response_text.strip()
        print(f"{section}: {formatted_response_text}")
        
        document.add_heading(section, level=1)
        document.add_paragraph(formatted_response_text)

        # break
        
    # Save the document to a temporary file
    document_path = 'response_document2.docx'
    document.save(document_path)
    return send_file(document_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')

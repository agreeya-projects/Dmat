from flask import Flask, request, jsonify, send_file
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain.chains import ConversationChain
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from flask_cors import CORS
from langchain_openai import OpenAI
import tiktoken
import json
from langchain_openai import ChatOpenAI
from langchain_mongodb import MongoDBAtlasVectorSearch
from openai import OpenAI as openai_OpenAI
import urllib.parse
import os
from dotenv import load_dotenv
from pymongo import MongoClient
from utils import set_open_key
from langchain_community.embeddings import HuggingFaceEmbeddings
from docx import Document as DocxDocument  
from langchain.llms import LlamaCpp
from langchain.docstore.document import Document
from groq import Groq
from rouge import Rouge
import re

# Load environment variables from .env file
load_dotenv()

# MongoDB credentials
username = urllib.parse.quote_plus(os.getenv('MONGO_USERNAME'))
password = urllib.parse.quote_plus(os.getenv('MONGO_PASSWORD'))
atlas_connection_string = f"mongodb+srv://{username}:{password}@cluster0.qualb0u.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"

# Connect to MongoDB
mongodb_client = MongoClient(atlas_connection_string)
print("-------------------------------------- MongoDB Connected ---------------------------------------------")
print(mongodb_client.list_database_names())

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes


client = Groq(
    api_key=os.environ.get("GROQ_API_KEY"),
)

db_name = "DMAT"


#The following is a friendly conversation between a human and an AI. The AI is talkative and provides lots of specific details from its context. If the AI does not know the answer to a question, it truthfully says it does not know.
#     The output content should not contain any extra text such as "as asked in query", "as an AI assistant", etc.
    # The output should be in a professional format suitable for direct inclusion in a documentation file. Exclude any extraneous text such as "as asked in query," "as an AI assistant," etc. 
    # The content should follow the given template and include all required sections and subsections.    

# Function to generate prompt based on the section
def generate_prompt(section_name, section_instructions, context_text):
    bos = "<|begin_of_text|><|start_header_id|>system<|end_header_id|>\n"
    user = "\n<|eot_id|><|start_header_id|>user<|end_header_id|>\n"
    assistant = "\n<|eot_id|><|start_header_id|>assistant<|end_header_id|>"
    base_system_message = """
    You are a helpful, smart, kind, and efficient AI assistant.
    You are a professional tool that generates documentation of any tool or script using the data provided in the database.
    """

    prompt_template = f"""
    {bos}
    Human: 
    {base_system_message}
    {user}
    
    Current conversation:
    Human:
    Generate detailed content for the {section_name} section of the documentation of the tool named DMAT based on the following instructions:
    {section_instructions}

    It is crucial that you use the information provided in the context to construct the response. Ensure that the content is deeply relevant to the detailed tool's features and functionalities mentioned in the context. Ensure that each section is comprehensive  engaging, and tailored specifically to the tool's features and functionalities.
    Avoid generic descriptions and focus on specifics related to the tool.

    The output should be professional and structured for direct inclusion in a documentation file. 
    Avoid any extra text such as "as asked in query," "as an AI assistant," etc. 
    Follow the provided template and ensure all required sections and subsections are included.
    Do not exceeds your response more than 300 words and strictly follow this in every sections.
    

    Provide the response in a structured and formal manner.
    {assistant}
    """
    if context_text:
        prompt_template += "\n\nContext:\n" + context_text
   
    return prompt_template


collection_name = "Information_Store"
# llm = LlamaCpp(
#                 #streaming = True,
#                 model_path="models/Meta-Llama-3-8B-Instruct.Q4_K_M.gguf",
#                 # config={#'max_new_tokens':128,
#                 #         'temperature':0.01,
#                 #         'lctx': 1500}
#                 temperature=0.8,
#                 # top_p=1,
#                 # verbose=True,
#                 n_ctx=4000
#                 )
# llm_agent = llm

# memory = ConversationBufferWindowMemory(llm=llm_agent, max_token_limit=4000)
# conversation = ConversationChain(llm=llm, memory=memory, verbose=True)

def get_vector_index_mongo_for_lang( collection_name, atlas_connection_string, db_name): #collection_name_suffix,
    
    embed_model = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5")  # Use HuggingFaceEmbeddings
    
    #collection_name = collection_name + "_" + collection_name_suffix
    vector_search = MongoDBAtlasVectorSearch.from_connection_string(
        atlas_connection_string,
        f"{db_name}.{collection_name}",
        embed_model,
        index_name=collection_name + "_index"
    )
    return vector_search

db = get_vector_index_mongo_for_lang(collection_name, atlas_connection_string, db_name) #collection_name_suffix, 
#print('indexes:-', db.collection.getIndexes())

# Example usage for multiple sections
sections = [
    {"name": "Introduction", "instructions": "For Introduction of the DMAT tool, generate a detailed paragraph highlighting its main features, functionalities, and use cases. Mention how it provides real-time monitoring and analysis for Android 5G NR and LTE devices. Include its compatibility, logging capabilities, Single-Sign-On authentication, and remote control features. Make the introduction informative and engaging for potential users and stakeholders. Please include any additional relevant information to make the introduction comprehensive and detailed."},
    {"name": "Purpose of the SDD", "instructions": "Write a comprehensive and detailed paragraph for the 'Purpose of the SDD' section. Explain the objective and significance of the Software Design Document (SDD) in the context of the DMAT tool. Describe how the SDD serves as a comprehensive blueprint for the software architecture, detailing specific design decisions and implementation strategies. Emphasize the importance of the SDD as a critical reference document for developers, testers, and stakeholders, ensuring alignment and coherence throughout the project lifecycle. Discuss how the SDD facilitates effective communication and understanding among all project participants, helping to identify and mitigate potential risks and challenges early in the process. Highlight the role of the SDD in maintaining consistency and clarity, ensuring that all aspects of the project are thoroughly documented and understood. Explain how the SDD supports continuous improvement and future scalability of the DMAT tool. Conclude by discussing how the SDD is indispensable for the successful implementation, deployment, and maintenance of the DMAT tool, ultimately contributing to its long-term success and efficiency. Please provide additional relevant details to ensure thorough coverage of the purpose."},
    {"name": "General Overview", "instructions": "Write a comprehensive and detailed paragraph for the 'General Overview' section of the DMAT tool. Describe the primary purpose of DMAT, highlighting its role in providing real-time monitoring and analysis for Android 5G NR and LTE devices. Explain how DMAT supports efficient device performance monitoring and facilitates various testing environments. Detail its advanced logging capabilities, emphasizing how they contribute to effective data collection and analysis. Discuss the tool's compatibility with a wide range of devices and the benefits of this broad compatibility. Highlight the Single-Sign-On authentication feature, explaining its role in enhancing security and simplifying user management. Describe the remote control features, including how they allow users to manage multiple devices simultaneously and improve operational efficiency. Mention the integration of DMAT with existing post-processing tools and the advantages this offers to Verizon's testing teams. Conclude by explaining how DMAT enhances overall productivity and efficiency, providing a robust and cost-effective solution for comprehensive device performance monitoring and analysis. Please include any additional relevant information to provide a thorough overview."},
    {"name": "Key Pain Points Addressed", "instructions": "Write a comprehensive and detailed paragraph for the 'Key Pain Points Addressed' section. Identify and describe each key pain point that the DMAT tool addresses. Include specific challenges related to real-time monitoring and analysis of Android 5G NR and LTE devices. Discuss how DMAT resolves issues such as inconsistent data collection, difficulty in managing multiple devices, and security concerns. Highlight the tool's advanced logging capabilities and how they improve data accuracy and reliability. Explain how the Single-Sign-On authentication feature enhances security and simplifies user management. Detail how the remote control features help in efficiently managing multiple devices simultaneously, reducing operational complexity. Mention the integration with existing post-processing tools and how it streamlines the analysis process for testing teams. Conclude by summarizing how DMAT effectively addresses these pain points, leading to improved efficiency, productivity, and overall performance in device monitoring and analysis. Please provide additional relevant details to ensure thorough coverage of key pain points."},    
    {"name": "Use Case Solution", "instructions": "Generate a detailed and comprehensive paragraph for the 'Use Case Solution' section. Describe specific scenarios where the DMAT tool is effectively used, highlighting real-time monitoring and analysis, data collection improvements, secure access via Single-Sign-On, remote management of multiple devices, and integration with post-processing tools. Provide examples of how DMAT enhances operational efficiency and reduces complexity in real-world applications."},
    {"name": "Goals and Expected Outcomes", "instructions": "Generate a detailed and comprehensive paragraph for the 'Goals and Expected Outcomes' section. Outline the primary goals of implementing the DMAT tool, such as improving real-time monitoring and analysis, enhancing data accuracy, and streamlining device management. Describe the expected outcomes, including increased operational efficiency, better security through Single-Sign-On, simplified user management, and more effective data integration with post-processing tools. Highlight the benefits to users and stakeholders, leading to improved performance and productivity."},
    {"name": "Scope of Proof of Concept (POC)", "instructions": "Generate a detailed and comprehensive paragraph for the 'Scope of Proof of Concept (POC)' section. Define the objectives and boundaries of the POC for the DMAT tool. Describe the specific features and functionalities to be tested, including real-time monitoring, advanced logging capabilities, Single-Sign-On authentication, and remote device management. Explain the criteria for success, performance benchmarks, user feedback, and integration with existing systems. Highlight the key deliverables and timeline for completing the POC, and discuss how the POC outcomes will inform the full-scale implementation and potential improvements for the DMAT tool."},
    {"name": "Assumptions", "instructions": "List and explain the assumptions made during the development of the DMAT tool."},
    {"name": "System Integration", "instructions": "Discuss the integration processes and procedures for the DMAT tool within the existing systems."},
    {"name": "Limitations / Out of Scope", "instructions": "Identify the limitations and out-of-scope elements of the DMAT tool."},
    {"name": "High Level Architecture and Design", "instructions": "Provide a high-level overview of the architecture and design of the DMAT tool."},
    {"name": "Security Procedures", "instructions": "Detail the security procedures implemented in the DMAT tool."},
    {"name": "Information Architecture – Data Flow", "instructions": "Describe the information architecture and data flow of the DMAT tool."},
    {"name": "Design and Usability", "instructions": "Discuss the design and usability aspects of the DMAT tool."},
    {"name": "System Design", "instructions": "Discuss the detailed system design of the DMAT tool."},
    {"name": "External Interfaces", "instructions": "Identify the external interfaces used by the DMAT tool."},
    {"name": "FAQs", "instructions": "List common questions and answers related to the DMAT tool."},
    {"name": "User Stories", "instructions": "List and describe the comprehensive user stories for the DMAT tool."},
    {"name": "Testing and Validation", "instructions": "Discuss the testing and validation processes for the DMAT tool."},
    {"name": "Deployment and Maintenance", "instructions": "Discuss the deployment and maintenance plans for the DMAT tool."},
    {"name": "Appendix", "instructions": "Include additional information related to the DMAT tool."}
]

toc_items = [
    "Introduction",
    "Purpose of the SDD",
    "General Overview",
    "Key Pain Points Addressed",
    "Use Case Solution",
    "Goals and Expected Outcomes",
    "Scope of Proof of Concept (POC)",
    "Assumptions",
    "System Integration",
    "Limitations / Out of Scope",
    "High Level Architecture and Design",
    "Security Procedures",
    "Information Architecture – Data Flow",
    "Design and Usability",
    "System Design",
    "External Interfaces",
    "FAQs",
    "User Stories",
    "Testing and Validation",
    "Deployment and Maintenance",
    "Appendix"
]



document = DocxDocument()
document.add_heading('Documentation of DMAT Tool', 0)
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def add_bold_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)  # Split text by **bold** sections
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]  # Remove ** from the start and end
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            paragraph.add_run(part)

def add_toc(document, toc_items):
    # Add a heading for the Table of Contents
    toc_heading = document.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for index, item in enumerate(toc_items, start=1):
        toc_entry = document.add_paragraph(f"{index}. {item}")
        toc_entry.style = 'Normal'
        toc_entry_format = toc_entry.paragraph_format
        toc_entry_format.left_indent = Pt(36)  # Set left indentation for TOC entries
        toc_entry_format.space_after = Pt(12)  # Add space after each TOC entry

# Add TOC to the document
add_toc(document, toc_items)

for section in sections:
    query_text = f"Generate detailed content for the {section['name']} section."
    results = db.similarity_search_with_score(query_text, k=3)
    context_text = "\n".join([doc.page_content for doc, score in results])
    prompt = generate_prompt(section["name"], section["instructions"], context_text)

    chat_completion = client.chat.completions.create(
    messages=[
        {
            "role": "user",
            "content": prompt,
        }
    ],
    model="llama-3.1-70b-versatile",
    )
    #response_text = conversation.predict(input=prompt)
    response_text = chat_completion.choices[0].message.content

    formatted_response_text = response_text.strip()

    #clean_response_text  = remove_bold_markdown(formatted_response_text)
    
    # rouge = Rouge()
    # scores = rouge.get_scores(formatted_response_text, context_text)
    # print(f"{section['name']} Response Rouge score with context", scores[0])
    #print(f"{section['name']}: {formatted_response_text}")

    paragraph = document.add_paragraph()
    add_bold_text(paragraph, formatted_response_text)
    
    # document.add_heading(section['name'], level=1)
    # document.add_paragraph(formatted_response_text)

# Save the document to a temporary file
document_path = 'Output_Document/Dmat_document_with_toc2.docx'
document.save(document_path)


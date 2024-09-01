from flask import Flask, request, jsonify, send_file
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain.chains import ConversationChain
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from flask_cors import CORS
from langchain_openai import OpenAI
import tiktoken
import json
from langchain_openai import ChatOpenAI
from langchain_mongodb import MongoDBAtlasVectorSearch
from openai import OpenAI as openai_OpenAI
import urllib.parse
import os
from dotenv import load_dotenv
from pymongo import MongoClient
from utils import set_open_key
from langchain_community.embeddings import HuggingFaceEmbeddings
from docx import Document as DocxDocument  
from langchain.llms import LlamaCpp
from langchain.docstore.document import Document
from groq import Groq
from rouge import Rouge
import re

# Load environment variables from .env file
load_dotenv()

# MongoDB credentials
username = urllib.parse.quote_plus(os.getenv('MONGO_USERNAME'))
password = urllib.parse.quote_plus(os.getenv('MONGO_PASSWORD'))
atlas_connection_string = f"mongodb+srv://{username}:{password}@cluster0.qualb0u.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"

# Connect to MongoDB
mongodb_client = MongoClient(atlas_connection_string)
print("-------------------------------------- MongoDB Connected ---------------------------------------------")
print(mongodb_client.list_database_names())

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes


client = Groq(
    api_key=os.environ.get("GROQ_API_KEY"),
)

db_name = "DMAT"


#The following is a friendly conversation between a human and an AI. The AI is talkative and provides lots of specific details from its context. If the AI does not know the answer to a question, it truthfully says it does not know.
#     The output content should not contain any extra text such as "as asked in query", "as an AI assistant", etc.
    # The output should be in a professional format suitable for direct inclusion in a documentation file. Exclude any extraneous text such as "as asked in query," "as an AI assistant," etc. 
    # The content should follow the given template and include all required sections and subsections.    

    #    It is crucial that you use the information provided in the context to construct the response. Ensure that the content is deeply relevant to the detailed tool's features and functionalities mentioned in the context. Ensure that each section is comprehensive  engaging, and tailored specifically to the tool's features and functionalities.
    # Avoid generic descriptions and focus on specifics related to the tool.

    # The output should be professional and structured for direct inclusion in a documentation file. 
    # Avoid any extra text such as "as asked in query," "as an AI assistant," etc. 
    # Follow the provided template and ensure all required sections and subsections are included.
    # Do not exceeds your response more than 300 words and strictly follow this in every sections.
    # Give me the detailed table of content.

# Function to generate prompt based on the section
def generate_prompt(section_name, section_instructions, context_text):
    bos = "<|begin_of_text|><|start_header_id|>system<|end_header_id|>\n"
    user = "\n<|eot_id|><|start_header_id|>user<|end_header_id|>\n"
    assistant = "\n<|eot_id|><|start_header_id|>assistant<|end_header_id|>"
    base_system_message = """
    You are a helpful, smart, kind, and efficient AI assistant.
    You are a professional tool that generates table of content of any tool or script.
    """

    prompt_template = f"""
    {bos}
    Human: 
    {base_system_message}
    {user}
    
    Current conversation:
    Human:
    Generate table of content for the {section_name} section of the documentation of the tool named DMAT based on the following instructions:
    {section_instructions}

 

    Provide the response in a structured and formal manner.
    {assistant}
    """
    if context_text:
        prompt_template += "\n\nContext:\n" + context_text
   
    return prompt_template


collection_name = "Information_Store"
# llm = LlamaCpp(
#                 #streaming = True,
#                 model_path="models/Meta-Llama-3-8B-Instruct.Q4_K_M.gguf",
#                 # config={#'max_new_tokens':128,
#                 #         'temperature':0.01,
#                 #         'lctx': 1500}
#                 temperature=0.8,
#                 # top_p=1,
#                 # verbose=True,
#                 n_ctx=4000
#                 )
# llm_agent = llm

# memory = ConversationBufferWindowMemory(llm=llm_agent, max_token_limit=4000)
# conversation = ConversationChain(llm=llm, memory=memory, verbose=True)

def get_vector_index_mongo_for_lang( collection_name, atlas_connection_string, db_name): #collection_name_suffix,
    
    embed_model = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5")  # Use HuggingFaceEmbeddings
    
    #collection_name = collection_name + "_" + collection_name_suffix
    vector_search = MongoDBAtlasVectorSearch.from_connection_string(
        atlas_connection_string,
        f"{db_name}.{collection_name}",
        embed_model,
        index_name=collection_name + "_index"
    )
    return vector_search

db = get_vector_index_mongo_for_lang(collection_name, atlas_connection_string, db_name) #collection_name_suffix, 
#print('indexes:-', db.collection.getIndexes())

# Example usage for multiple sections
sections = [
    {"name": "Introduction", "instructions": "Generate only the title for the 'Introduction' section to be used in the table of contents."},
    {"name": "Purpose of the SDD", "instructions": "Generate only the title for the 'Purpose of the SDD' section to be used in the table of contents."},
    {"name": "General Overview", "instructions": "Generate only the title for the 'General Overview' section to be used in the table of contents."},
    {"name": "Key Pain Points Addressed", "instructions": "Generate only the title for the 'Key Pain Points Addressed' section to be used in the table of contents."},
    {"name": "Use Case Solution", "instructions": "Generate only the title for the 'Use Case Solution' section to be used in the table of contents."},
    {"name": "Goals and Expected Outcomes", "instructions": "Generate only the title for the 'Goals and Expected Outcomes' section to be used in the table of contents."},
    {"name": "Scope of Proof of Concept (POC)", "instructions": "Generate only the title for the 'Scope of Proof of Concept (POC)' section to be used in the table of contents."},
    {"name": "Assumptions", "instructions": "Generate only the title for the 'Assumptions' section to be used in the table of contents."},
    {"name": "System Integration", "instructions": "Generate only the title for the 'System Integration' section to be used in the table of contents."},
    {"name": "Limitations / Out of Scope", "instructions": "Generate only the title for the 'Limitations / Out of Scope' section to be used in the table of contents."},
    {"name": "High Level Architecture and Design", "instructions": "Generate only the title for the 'High Level Architecture and Design' section to be used in the table of contents."},
    {"name": "Security Procedures", "instructions": "Generate only the title for the 'Security Procedures' section to be used in the table of contents."},
    {"name": "Information Architecture – Data Flow", "instructions": "Generate only the title for the 'Information Architecture – Data Flow' section to be used in the table of contents."},
    {"name": "Design and Usability", "instructions": "Generate only the title for the 'Design and Usability' section to be used in the table of contents."},
    {"name": "System Design", "instructions": "Generate only the title for the 'System Design' section to be used in the table of contents."},
    {"name": "External Interfaces", "instructions": "Generate only the title for the 'External Interfaces' section to be used in the table of contents."},
    {"name": "FAQs", "instructions": "Generate only the title for the 'FAQs' section to be used in the table of contents."},
    {"name": "User Stories", "instructions": "Generate only the title for the 'User Stories' section to be used in the table of contents."},
    {"name": "Testing and Validation", "instructions": "Generate only the title for the 'Testing and Validation' section to be used in the table of contents."},
    {"name": "Deployment and Maintenance", "instructions": "Generate only the title for the 'Deployment and Maintenance' section to be used in the table of contents."},
    {"name": "Appendix", "instructions": "Generate only the title for the 'Appendix' section to be used in the table of contents."}
]



document = DocxDocument()
document.add_heading('Documentation', 0)
from docx.shared import Pt
def add_bold_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)  # Split text by **bold** sections
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]  # Remove ** from the start and end
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            paragraph.add_run(part)

for section in sections:
    query_text = f"Generate table of content for the {section['name']} section According to the instructions mentioned for this section."
    results = db.similarity_search_with_score(query_text, k=3)
    context_text = "\n".join([doc.page_content for doc, score in results])
    prompt = generate_prompt(section["name"], section["instructions"], context_text)

    chat_completion = client.chat.completions.create(
    messages=[
        {
            "role": "user",
            "content": prompt,
        }
    ],
    model="llama-3.1-70b-versatile",
    )
    #response_text = conversation.predict(input=prompt)
    response_text = chat_completion.choices[0].message.content

    formatted_response_text = response_text.strip()

    #clean_response_text  = remove_bold_markdown(formatted_response_text)
    
    # rouge = Rouge()
    # scores = rouge.get_scores(formatted_response_text, context_text)
    # print(f"{section['name']} Response Rouge score with context", scores[0])
    #print(f"{section['name']}: {formatted_response_text}")

    paragraph = document.add_paragraph()
    add_bold_text(paragraph, formatted_response_text)
    
    # document.add_heading(section['name'], level=1)
    # document.add_paragraph(formatted_response_text)

# Save the document to a temporary file
document_path = 'table_of_content1.docx'
document.save(document_path)


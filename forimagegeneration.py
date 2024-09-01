import os
import pdfplumber
import openai
from sentence_transformers import SentenceTransformer, util
from PIL import Image, ImageOps
import imagehash
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import time
import warnings
import cv2
from scipy.spatial.distance import hamming
from dotenv import load_dotenv
from docx import Document

# Load environment variables
load_dotenv()

# Suppress specific warnings
warnings.filterwarnings("ignore", category=FutureWarning, module="huggingface_hub.file_download")

# Set your OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Extract images from PDF using pdfplumber
def extract_images_from_pdf(pdf_path, output_folder):
    images = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                for img in page.images:
                    if img["width"] > 100 and img["height"] > 100:
                        x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
                        x0 = max(x0, 0)
                        y0 = max(y0, 0)
                        x1 = min(x1, page.width)
                        y1 = min(y1, page.height)
                        cropped_image = page.within_bbox((x0, y0, x1, y1)).to_image(resolution=900)  # Increase resolution
                        pil_img = cropped_image.original
                        pil_img = pil_img.resize((int(pil_img.width * 2), int(pil_img.height * 2)), Image.LANCZOS)
                        image_path = os.path.join(output_folder, f"{os.path.basename(pdf_path)}_page_{i + 1}_image_{len(images) + 1}.png")
                        pil_img.save(image_path, format='PNG', quality=95)
                        images.append(image_path)
    except Exception as e:
        print(f"Error processing {pdf_path}: {e}")
    return images

# Get embeddings for texts
def get_embeddings(texts):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    embeddings = model.encode(texts, convert_to_tensor=True)
    return embeddings

# Separate function to capture screenshot using Selenium
def capture_screenshot(url, output_path):
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--window-size=1920,1080")
    driver = webdriver.Chrome(options=chrome_options)
    driver.get(url)
    time.sleep(5)
    driver.save_screenshot(output_path)
    driver.quit()

# Capture frames from video
def capture_frames_from_video(video_path, output_folder, interval=3):
    video_capture = cv2.VideoCapture(video_path)
    fps = int(video_capture.get(cv2.CAP_PROP_FPS))
    success, image = video_capture.read()
    count = 0
    while success:
        if count % (fps * interval) == 0:
            frame_path = os.path.join(output_folder, f"frame_{count}.png")
            cv2.imwrite(frame_path, image)
        success, image = video_capture.read()
        count += 1
    video_capture.release()

# Remove duplicate and similar images
def remove_duplicate_images(image_paths, threshold=0.9):
    unique_images = []
    unique_hashes = []

    for img_path in image_paths:
        try:
            img = Image.open(img_path)
            img = ImageOps.exif_transpose(img)  # Handle orientation based on EXIF
            img_hash = imagehash.phash(img)

            if not any(hamming(img_hash.hash, u_hash.hash) < threshold for u_hash in unique_hashes):
                unique_images.append(img_path)
                unique_hashes.append(img_hash)
            else:
                os.remove(img_path)  # Remove duplicate image
        except Exception as e:
            print(f"Error processing image {img_path}: {e}")

    return unique_images

# Generate descriptions for images
def generate_image_description(image_path, context):
    with open(image_path, "rb") as image_file:
        image_data = image_file.read()

    prompt = f"Provide a detailed description for the following image in the context of {context}."
    
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an AI assistant."},
            {"role": "user", "content": prompt},
        ],
    )
    
    description = response.choices[0].message["content"]
    return description

# Save descriptions to a document
def save_descriptions_to_doc(image_descriptions, doc_path):
    doc = Document()
    doc.add_heading('Image Descriptions', level=1)
    
    for image_name, description in image_descriptions.items():
        doc.add_heading(image_name, level=2)
        doc.add_paragraph(description)
    
    doc.save(doc_path)

# Paths
pdf_folder = 'all_data/pdf'
video_folder = 'all_data/video'
output_folder = 'all_data/images'
doc_path = 'image_descriptions.docx'
query = "Provide detailed information of High-Level Design (HLD) and Low-Level Design (LLD). HLD includes the main components of the system, their interactions, and architecture overview. LLD will include specific modules, components, their interactions, data flow, and implementation details."

# Ensure output folder exists
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Extract images from PDFs
pdf_images = {}
for pdf_file in os.listdir(pdf_folder):
    if pdf_file.endswith('.pdf'):
        pdf_path = os.path.join(pdf_folder, pdf_file)
        if os.path.exists(pdf_path):
            images = extract_images_from_pdf(pdf_path, output_folder)
            pdf_images[pdf_file] = images
        else:
            print(f"File not found: {pdf_path}")

# Get query embedding
query_embedding = get_embeddings([query])

# Get embeddings for PDFs
pdf_texts = {pdf_file: " ".join(pdf_images[pdf_file]) for pdf_file in pdf_images}
pdf_embeddings = {pdf_file: get_embeddings([text]) for pdf_file, text in pdf_texts.items()}

# Calculate similarity scores
similarity_scores = {}
for pdf_file, embedding in pdf_embeddings.items():
    similarity = util.pytorch_cos_sim(query_embedding, embedding).item()
    similarity_scores[pdf_file] = similarity

# Get top 3 similar PDFs
top_pdfs = sorted(similarity_scores, key=similarity_scores.get, reverse=True)[:3]

# Capture screenshots of website pages
pages = {
    'login_page': "https://ilogin.verizon.com/ngauth/verifyusercontroller?goto=https://ssologin.verizon.com/ngauth/idpssoinit?metaAlias%3D/employee/VZIDP_24626_CKVV_DMATFED%26spEntityID%3DDMATPRODAWS",
    'home_page': "https://dmat.verizon.com/#/home",
    'manage_password': "https://ilogin.verizon.com/vzsplogin/LoginController?TYPE=33554433&REALMOID=06-d7da31f1-a37a-108d-bd4b-85115b4afd5f&GUID=&SMAUTHREASON=0&METHOD=GET&SMAGENTNAME=$SM$9%2f4gVFwW6gsaabjchgJmA3wGd97AYDbqZtRSRHob8tnsTNpMHiMKXB4NJ9aFW%2bix&TARGET=$SM$https%3a%2f%2fvzsp%2everizon%2ecom%2fvzsp%2fhomecontroller",
    'find_my_id': "https://ilogin.verizon.com/manageaccount/FindMyIDController",
    'manage_profile': "https://ilogin.verizon.com/vzsplogin/LoginController?TYPE=33554433&REALMOID=06-d7da31f1-a37a-108d-bd4b-85115b4afd5f&GUID=&SMAUTHREASON=0&METHOD=GET&SMAGENTNAME=$SM$9%2f4gVFwW6gsaabjchgJmA3wGd97AYDbqZtRSRHob8tnsTNpMHiMKXB4NJ9aFW%2bix&TARGET=$SM$https%3a%2f%2fvzsp%2everizon%2ecom%2fvzsp%2fhomecontroller"
}

for page, url in pages.items():
    capture_screenshot(url, os.path.join(output_folder, f'{page}.png'))

# Extract frames from videos
for video_file in os.listdir(video_folder):
    if video_file.endswith(('.mp4', '.avi', '.mkv')):
        video_path = os.path.join(video_folder, video_file)
        capture_frames_from_video(video_path, output_folder)

# Collect all images for de-duplication
all_image_paths = []
for img_list in pdf_images.values():
    all_image_paths.extend(img_list)

for page in pages.keys():
    all_image_paths.append(os.path.join(output_folder, f'{page}.png'))

for img_file in os.listdir(output_folder):
    if img_file.startswith('frame_'):
        all_image_paths.append(os.path.join(output_folder, img_file))

# Remove duplicate and similar images
unique_image_paths = remove_duplicate_images(all_image_paths)

# Generate descriptions for unique images
image_descriptions = {}
for img_path in unique_image_paths:
    description = generate_image_description(img_path, query)
    image_name = os.path.basename(img_path)
    image_descriptions[image_name] = description

# Save descriptions to a document
save_descriptions_to_doc(image_descriptions, doc_path)

print("Unique images saved and descriptions written to document.")
for img_path in unique_image_paths:
    print(img_path)

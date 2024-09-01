import os
import pymongo
import re
import urllib.parse
import openai
import pandas as pd
import xml.etree.ElementTree as ET
import yaml
from docx import Document as DocxDocument
from dotenv import load_dotenv
from llama_index.core import StorageContext, Document, VectorStoreIndex, SimpleDirectoryReader
from llama_index.vector_stores.mongodb import MongoDBAtlasVectorSearch
from llama_index.core import Settings
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.readers.file import CSVReader
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from vid2aud import VideoToaudio_converter
#from aud2tex import audio2text

# Load environment variables
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# MongoDB credentials
username = urllib.parse.quote_plus(os.getenv('MONGO_USERNAME'))
password = urllib.parse.quote_plus(os.getenv('MONGO_PASSWORD'))
atlas_connection_string = f"mongodb+srv://{username}:{password}@cluster0.qualb0u.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"

# Connect to MongoDB
mongodb_client = pymongo.MongoClient(atlas_connection_string)

# Database and collection names
db_name = "DMAT"
collection_name = "Information_Store"

# Check if database exists
if db_name in mongodb_client.list_database_names():
    db = mongodb_client[db_name]
else:
    db = mongodb_client[db_name]  # This will create the database if it does not exist

# Check if collection exists
if collection_name not in db.list_collection_names():
    db.create_collection(collection_name)  # This will create the collection if it does not exist

print(f"Database '{db_name}' and collection '{collection_name}' are ready.")

required_exts = [".md", ".txt", ".pdf", ".PDF", ".docx", ".DOCX", ".xlsx", ".xml", ".yaml"]

headers_to_split_on = [
    ("#", "title"),
    ("##", "header"),
    ("###", "header"),
    ("####", "header"),
    ("#####", "header"),
    ("######", "header"),
]

markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=2000,
    chunk_overlap=500,
    length_function=len,
    add_start_index=True,
)

def clean_text(text):
    text = re.sub(r'\b\w{101,}\b', ' ', text)
    text = re.sub(r' +', ' ', text)
    return text

def save_to_mongo(chunks: list[Document], db_name, collection_name, local_embedding):
    try:
        mongodb_client = pymongo.MongoClient(atlas_connection_string, tls=True, tlsAllowInvalidCertificates=True)
        index_name = collection_name + "_index"
        atlas_vector_search = MongoDBAtlasVectorSearch(
            mongodb_client,
            db_name=db_name,
            collection_name=collection_name,
            index_name=index_name
        )
        
        if local_embedding:
            embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en-v1.5")
            Settings.embed_model = embed_model

        vector_store_context = StorageContext.from_defaults(vector_store=atlas_vector_search)
        vector_store_index = VectorStoreIndex.from_documents(
            chunks, storage_context=vector_store_context, show_progress=True
        )

        return vector_store_index
    except pymongo.errors.ServerSelectionTimeoutError as e:
        print("Server selection timeout error:", e)
    except pymongo.errors.ConnectionError as e:
        print("Connection error:", e)
    except Exception as e:
        print("An error occurred while connecting to MongoDB:", e)

def generate_data_store_mongo(output_file_path, local_embedding):
    documents = load_files(output_file_path, required_exts)
    save_to_mongo(documents, db_name, collection_name, local_embedding)

def push_csv(url_name, local_embedding):
    parser = CSVReader()
    file_extractor = {".csv": parser}
    documents = SimpleDirectoryReader(
        os.path.join("tempDataFolder", "csv", url_name), file_extractor=file_extractor
    ).load_data()
    
    for counter, i in enumerate(documents):
        if len(".".join(i.metadata["filename"].split(".")[:-1]).split("_")) == 1:
            documents[counter].metadata["id"] = ".".join(i.metadata["filename"].split(".")[:-1]).split("_")[0]
        else:
            documents[counter].metadata["id"] = "_".join(".".join(i.metadata["filename"].split(".")[:-1]).split("_")[:-1])
    
    save_to_mongo(documents, db_name, collection_name, local_embedding)

def load_files(input_path, required_exts):
    documents = []
    filename_fn = lambda filename: {"file_name": filename.split("/")[-1]}
    
    if input_path.endswith('.xlsx'):
        txt_path = convert_xlsx_to_txt(input_path)
        documents.extend(load_txt(txt_path))
        os.remove(txt_path)
    elif input_path.endswith('.yaml') or input_path.endswith('.yml'):
        txt_path = convert_yaml_to_txt(input_path)
        if txt_path:
            documents.extend(load_txt(txt_path))
            os.remove(txt_path)
    else:
        if os.path.isdir(input_path):
            reader = SimpleDirectoryReader(
                input_dir=input_path,
                required_exts=required_exts,
                file_metadata=filename_fn,
                filename_as_id=True,
                recursive=True,
            )
        else:
            reader = SimpleDirectoryReader(
                input_files=[input_path],
                required_exts=required_exts,
                file_metadata=filename_fn,
                filename_as_id=True,
                recursive=True,
            )

        md_docs = reader.load_data()

        for md_doc in md_docs:
            md_doc.text = clean_text(md_doc.text)
            docs = markdown_splitter.split_text(md_doc.text)
            docs = text_splitter.split_documents(docs)
            
            for doc in docs:
                _tmp_path = ".".join(md_doc.doc_id.split(os.path.sep)[-1].split("_"))
                doc.metadata["id"] = _tmp_path
                documents.append(Document(
                    text=doc.page_content,
                    metadata=doc.metadata,
                    excluded_llm_metadata_keys=["link"],
                    excluded_embed_metadata_keys=["link"]
                ))
    
    return documents

def convert_xlsx_to_txt(file_path):
    df = pd.read_excel(file_path)
    txt_content = df.to_csv(sep=' ', index=False)
    txt_path = file_path.replace('.xlsx', '.txt')
    with open(txt_path, 'w') as f:
        f.write(txt_content)
    return txt_path

def convert_yaml_to_txt(file_path):
    try:
        with open(file_path, 'r') as file:
            content = yaml.safe_load(file)
        txt_content = yaml.dump(content)
        txt_path = file_path.replace('.yaml', '.txt').replace('.yml', '.txt')
        with open(txt_path, 'w') as f:
            f.write(txt_content)
        return txt_path
    except yaml.YAMLError as e:
        print(f"YAML parsing error in file {file_path}: {e}")
        return None

def parse_xml(file_path):
    documents = []
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        text = ET.tostring(root, encoding='utf8', method='text').decode('utf8')
        documents.append(Document(text=clean_text(text), metadata={"file_name": os.path.basename(file_path)}))
    except Exception as e:
        print(f"Failed to parse XML file {file_path}: {e}")
    return documents

def load_txt(input_txt_path):
    documents = []
    filename_fn = lambda filename: {"file_name": filename.split("/")[-1]}
    
    if os.path.isdir(input_txt_path):
        reader = SimpleDirectoryReader(
            input_dir=input_txt_path,
            required_exts=[".txt"],
            file_metadata=filename_fn,
            filename_as_id=True,
            recursive=True,
        )
    else:
        reader = SimpleDirectoryReader(
            input_files=[input_txt_path],
            required_exts=[".txt"],
            file_metadata=filename_fn,
            filename_as_id=True,
            recursive=True,
        )

    txt_docs = reader.load_data()

    for txt_doc in txt_docs:
        txt_doc.text = clean_text(txt_doc.text)
        docs = text_splitter.split_documents([txt_doc])
        
        for doc in docs:
            _tmp_path = ".".join(txt_doc.doc_id.split(os.path.sep)[-1].split("_"))
            doc.metadata["id"] = _tmp_path
            documents.append(Document(
                text=doc.page_content,
                metadata=doc.metadata,
                excluded_llm_metadata_keys=["link"],
                excluded_embed_metadata_keys=["link"]
            ))
    
    return documents

def process_directory(root_directory, local_embedding):
    for root, _, files in os.walk(root_directory):
        for file in files:
            if any(file.endswith(ext) for ext in required_exts):
                file_path = os.path.join(root, file)
                try:
                    # Convert YAML and XLSX files to text before processing
                    if file_path.endswith('.xlsx'):
                        txt_path = convert_xlsx_to_txt(file_path)
                        generate_data_store_mongo(txt_path, local_embedding)
                        os.remove(txt_path)
                    elif file_path.endswith('.yaml') or file_path.endswith('.yml'):
                        txt_path = convert_yaml_to_txt(file_path)
                        if txt_path:
                            generate_data_store_mongo(txt_path, local_embedding)
                            os.remove(txt_path)
                    else:
                        generate_data_store_mongo(file_path, local_embedding)
                    print("Processed:", file_path)
                except Exception as e:
                    print(f"Failed to process file {file_path}: {e}")

#VideoToaudio_converter()
#audio2text()
# Example usage
root_directory = "all_data"
local_embedding = True

process_directory(root_directory, local_embedding)
